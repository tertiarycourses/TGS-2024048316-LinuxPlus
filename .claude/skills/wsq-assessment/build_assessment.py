#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Application Integration with Docker and Kubernetes' (TGS-2021010366):
  - Written Assessment (SAQ)  — 5 open-ended KNOWLEDGE questions (K1–K5), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (LO1–LO4), aligned to the in-class activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "CompTIA Certified Linux+ Training"
COURSE_CODE = "TGS-2024048316"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("comptia-linux-logo.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v1", "v1"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 ("K1",
  "A Linux server must be brought up from power-on and its files organised in a predictable way. Domain 1 "
  "(System Management) covers the boot chain and the Filesystem Hierarchy Standard.",
  "Describe the Linux boot process from power-on to a login prompt, explaining the roles of the firmware/UEFI, "
  "GRUB2, the kernel, the initramfs and systemd. What is the Filesystem Hierarchy Standard (FHS) and why does "
  "it matter to an administrator?",
  ["Firmware/UEFI (or BIOS) initialises the hardware and hands off to the boot loader.",
   "GRUB2 loads the Linux kernel and the initramfs, passing the kernel command line (root device, options).",
   "The kernel mounts the initramfs (a small cpio archive) as a temporary root to load storage/crypto modules, "
   "then mounts the real root filesystem.",
   "systemd (PID 1) starts services and reaches a target (e.g. multi-user.target / graphical.target), ending at "
   "a login prompt.",
   "The FHS is the standard layout of the root filesystem — /etc (config), /var (variable data/logs), /usr "
   "(programs), /home (users), /bin, /sbin, /proc & /sys (virtual). It matters because it lets admins and tools "
   "find files predictably across distributions. "
   "(Slides: Domain 1 — Key Concepts / Lab 1 — Boot Process & Filesystem Hierarchy)"]),
 ("K2",
  "Domain 2 (Services and User Management) covers how Linux stores local accounts and how work is run as "
  "services or on a schedule.",
  "Explain how Linux stores and manages local user accounts, including the roles of /etc/passwd, /etc/shadow "
  "and /etc/group. Then contrast a systemd service with a cron (or at) job for running work on a schedule.",
  ["/etc/passwd holds one line per account (username, UID, GID, home, login shell); it is world-readable.",
   "/etc/shadow holds the hashed passwords and password-ageing fields, readable only by root.",
   "/etc/group defines groups and their members; a user has one primary group and any number of supplementary "
   "groups (usermod -aG). Accounts are managed with useradd/adduser, usermod, passwd and chage.",
   "A systemd service (unit) runs and supervises a long-running or one-shot program — start/enable/status via "
   "systemctl, with logs in journalctl; timers can trigger units on a schedule.",
   "cron runs recurring five-field jobs (crontab), anacron catches missed jobs, and at runs a one-off job at a "
   "given time. Use a service/timer for supervised, journald-integrated tasks; use cron/at for simple scheduled "
   "commands. (Slides: Domain 2 — Key Concepts / Lab 9 — Accounts / Lab 10 — Processes & Scheduling / Lab 12 — systemd)"]),
 ("K3",
  "Domain 3 (Security) covers firewalling at several layers of the netfilter stack and operating-system "
  "hardening.",
  "Linux firewalling can be configured at several layers. Describe ufw, firewalld and nftables/iptables and "
  "explain what stateful (connection-tracking) inspection means. Give two OS-hardening techniques you would "
  "apply to a newly built server.",
  ["ufw (Uncomplicated Firewall) is a simple front end for allow/deny rules on ports/services — good for hosts.",
   "firewalld manages the firewall with zones and rich rules and is common on RHEL-family systems.",
   "iptables/nftables are the low-level netfilter front ends; nftables is the modern replacement and its inet "
   "family covers both IPv4 and IPv6. All of them ultimately program the kernel's netfilter framework.",
   "Stateful inspection means the firewall tracks connection state (NEW / ESTABLISHED / RELATED) via conntrack, "
   "so it can allow return traffic for permitted connections and only accept genuinely new sessions.",
   "OS-hardening examples (any two): harden SSH (PermitRootLogin no, key-only auth), apply least-privilege file "
   "permissions/ACLs and remove unneeded SUID bits, enforce SELinux/AppArmor, run fail2ban, disable cleartext "
   "services (telnet/FTP), keep the system patched. "
   "(Slides: Domain 3 — Key Concepts / Lab 15 — Firewalls / Lab 16 — OS Hardening)"]),
 ("K4",
  "Domain 4 (Automation, Orchestration and Scripting) covers Infrastructure as Code and shell scripting.",
  "What does it mean for an Ansible playbook to be idempotent, and why is that valuable? Then describe the core "
  "Bash scripting constructs you would use to automate an administrative task, and name one tool you would use "
  "to check a shell script for errors.",
  ["Idempotent means running the playbook repeatedly converges the system to the same desired state without "
   "making unnecessary changes — a task only changes something if it is not already correct.",
   "It is valuable because runs are safe to repeat, produce predictable results, and avoid drift; handlers fire "
   "only when a task reports 'changed', so services are not needlessly restarted.",
   "Core Bash constructs: a shebang (#!/usr/bin/env bash); variables and parameter expansion (\"${1:-default}\"); "
   "conditionals (if / case with [[ ]] test operators); loops (for / while / until); functions; and exit codes "
   "($?).",
   "Defensive scripts use set -euo pipefail and quote variables.",
   "shellcheck is the standard static analyser used to catch common shell bugs (unquoted variables, wrong test "
   "syntax). (Slides: Domain 4 — Key Concepts / Lab 20 — Ansible / Lab 21 — Bash Scripting)"]),
 ("K5",
  "Domain 5 (Troubleshooting) covers monitoring vocabulary and a repeatable method for diagnosing performance "
  "problems.",
  "Describe the universal performance-troubleshooting loop. For a CPU bottleneck, a memory-pressure problem and "
  "a disk-I/O bottleneck, name one command or counter you would use to confirm each. Briefly contrast SLA, SLI "
  "and SLO.",
  ["The loop: establish a baseline → reproduce/observe the spike → identify the bottleneck with the right "
   "counter → remediate → verify recovery against the baseline.",
   "CPU bottleneck: high load average vs core count, confirmed per-CPU with mpstat -P ALL or top (%us/%sy).",
   "Memory pressure: free -m shows low available memory and vmstat shows non-zero si/so (swap in/out).",
   "Disk-I/O bottleneck: iostat -x shows high %util and await (and processes stuck in D state); pidstat -d "
   "attributes the I/O.",
   "SLI = a measured indicator (e.g. request latency, error rate); SLO = the internal target for an SLI (e.g. "
   "99.9% of requests < 200 ms); SLA = the external, contractual promise (often with penalties) built on SLOs. "
   "(Slides: Domain 5 — Key Concepts / Lab 25 — Monitoring / Lab 29 — Performance)"]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "You have been hired as a junior Linux administrator at CompTech Solutions, a medium-sized tech company that "
 "has recently migrated its infrastructure to Linux to improve security and scalability. Your supervisor has "
 "tasked you with keeping the company's CRM application running smoothly: evaluating change requests, producing "
 "a short training guide, troubleshooting application issues from logs, analysing system performance, and "
 "deploying a new feature package. Complete the four tasks below; each mirrors a hands-on lab you did in class. "
 "For each task, paste your commands, your script and output snapshots, and any text-file contents as evidence.")

BOX_CAP = "Paste your commands, script and output snapshots, and text-file contents in the box below"
PRACTICAL = [
 ("Task 1", "LO4",
  "Evaluate change requests and write a training guide. You've been given several change requests for the CRM "
  "application — real-time data synchronisation, a new reporting module, and a database upgrade. "
  "Part A — Write a Bash script (evaluate_changes.sh) that inspects the host's resources (CPU cores, available "
  "memory and free disk space) and, using conditional logic and a loop, prints which requests are feasible on "
  "the current system. Use a shebang, set -euo pipefail, variables, an if/[[ ]] test and a for loop, then "
  "confirm the script is clean with shellcheck. "
  "Part B — Write a short training guide (crm_training.txt) covering how to install and maintain the CRM "
  "application. Provide your script, the output of running it, and the first lines of the training guide. "
  "(Aligns to Lab 21 — Bash Scripting; Lab 24 — Responsible AI / documentation.)",
  BOX_CAP,
  "Part A — evaluate_changes.sh:\n"
  "#!/usr/bin/env bash\n"
  "set -euo pipefail\n"
  "cores=$(nproc)\n"
  "mem_mb=$(free -m | awk '/^Mem:/{print $7}')            # available RAM (MB)\n"
  "disk_mb=$(df -m --output=avail / | tail -1 | tr -d ' ')\n"
  "# name:cores:mem_mb:disk_mb required\n"
  "requests=(\"realtime-sync:2:1024:5000\" \"reporting:1:512:2000\" \"db-upgrade:4:2048:10000\")\n"
  "for r in \"${requests[@]}\"; do\n"
  "  IFS=':' read -r name c m d <<<\"$r\"\n"
  "  if [[ $cores -ge $c && $mem_mb -ge $m && $disk_mb -ge $d ]]; then\n"
  "    echo \"FEASIBLE: $name\"\n"
  "  else\n"
  "    echo \"NOT FEASIBLE: $name (needs ${c} cores, ${m}MB RAM, ${d}MB disk)\"\n"
  "  fi\n"
  "done\n"
  "# Run & lint:\n"
  "chmod +x evaluate_changes.sh && ./evaluate_changes.sh\n"
  "shellcheck evaluate_changes.sh          # clean report\n\n"
  "Part B — crm_training.txt (first lines):\n"
  "CRM Application - Install & Maintenance Guide\n"
  "1. Install:  sudo apt install ./crm_analytics.deb\n"
  "2. Start:    sudo systemctl enable --now crm\n"
  "3. Logs:     journalctl -u crm -f\n\n"
  "Award the mark where the candidate uses resource checks (nproc / free / df), a conditional plus a loop, a "
  "shellcheck-clean script, and a clear training guide. (Lab 21 — Bash Scripting; Lab 24 — documentation.)"),
 ("Task 2", "LO5",
  "Troubleshoot the CRM from its logs. The CRM application has been experiencing slowdowns. Analyse the provided "
  "log file crm_app.log using Linux text-processing tools to find the errors and patterns behind the slowdown. "
  "Filter the log for errors, count how often each error type occurs, and identify the busiest time window. "
  "Document your findings and the steps you would take to resolve them in findings.txt. Provide the commands you "
  "used, the output showing the relevant errors, and your findings file. "
  "(Aligns to Lab 5 — Shell Operations & Text Processing; Lab 25 — Monitoring / log aggregation.)",
  BOX_CAP,
  "Filter and quantify the errors:\n"
  "grep -iE \"error|fail|timeout\" crm_app.log | head\n"
  "grep -c -i error crm_app.log                      # total error count\n"
  "awk '/ERROR/{print $NF}' crm_app.log | sort | uniq -c | sort -rn   # top error types\n"
  "# Busiest window (group by hour):\n"
  "awk '{print $2}' crm_app.log | cut -d: -f1 | sort | uniq -c | sort -rn | head\n"
  "# For a running service, the journald equivalent:\n"
  "journalctl -u crm --since \"1 hour ago\" -p err --no-pager\n\n"
  "findings.txt:\n"
  "- Most frequent error: database connection timeout (NN occurrences).\n"
  "- Peak error volume between 14:00-15:00.\n"
  "- Proposed fix: increase the DB connection-pool size, add the missing index, then restart crm.\n\n"
  "Award the mark for correct grep / awk / sort / uniq filtering, identifying the dominant error and the peak "
  "window, and a clear findings file with remediation steps. (Lab 5 — Text Processing; Lab 25 — Monitoring.)"),
 ("Task 3", "LO5",
  "Analyse performance and optimise. Review the provided performance_report.txt for the CRM host and identify at "
  "least three underlying issues contributing to poor performance (consider CPU, memory, disk I/O and "
  "processes). Then read user_feedback.txt and propose concrete changes to optimise the application. List the "
  "commands you would run on a live host to confirm each issue, write the three issues and your analysis to "
  "perf_issues.txt, and record your proposed changes. "
  "(Aligns to Lab 29 — Performance Troubleshooting; Lab 10 — Processes, Jobs & Scheduling.)",
  BOX_CAP,
  "Confirm each issue on a live host:\n"
  "uptime; mpstat -P ALL 1 3               # CPU: load > cores, a core pegged at 100%\n"
  "free -m; vmstat 1 5                      # Memory: low available; si/so > 0 = swapping\n"
  "iostat -x 1 3; pidstat -d 1 3            # Disk I/O: high %util/await; top writer PID\n"
  "ps -eo pid,pcpu,pmem,stat,cmd --sort=-pcpu | head   # top consumers; D = I/O wait\n\n"
  "perf_issues.txt (at least three):\n"
  "1. CPU saturation - load average exceeds the core count (mpstat shows a thread at 100%).\n"
  "2. Memory pressure - free shows little available RAM; vmstat si/so are non-zero (swapping).\n"
  "3. Disk I/O bottleneck - iostat %util near 100% with high await; a process stuck in D state.\n\n"
  "Proposed changes (from user_feedback.txt):\n"
  "- Add an index/cache to cut database CPU; fix the memory leak or add RAM; move the DB to faster storage; "
  "renice the heavy batch job so it stops starving the app.\n\n"
  "Award the mark for mapping each symptom to the right counter (load vs mpstat, free vs vmstat si/so, iostat "
  "%util/await), three genuine issues, and feedback-driven optimisations. (Lab 29 — Performance; Lab 10 — Processes.)"),
 ("Task 4", "LO2",
  "Deploy and test a new feature package. A new data-analytics feature is shipped as a Debian package, "
  "crm_analytics.deb. Install it, configure it to run as a service, and test it. Show the commands to install "
  "the package (resolving dependencies), enable and start it under systemd, confirm it is active and listening, "
  "and read its logs. After testing, write a short performance assessment plus one suggested enhancement to "
  "analytics_assessment.txt. Provide the install commands and output, the configure/test commands, and your "
  "assessment file. (Aligns to Lab 11 — Software & Package Management; Lab 12 — Service Management with systemd.)",
  BOX_CAP,
  "Install the .deb (apt resolves dependencies):\n"
  "sudo apt install ./crm_analytics.deb     # or: sudo dpkg -i crm_analytics.deb; sudo apt -f install\n"
  "dpkg -l | grep crm-analytics             # confirm installed\n"
  "dpkg -L crm-analytics | head             # files the package placed\n\n"
  "Configure and run it as a service:\n"
  "sudo systemctl daemon-reload\n"
  "sudo systemctl enable --now crm-analytics\n"
  "systemctl status crm-analytics --no-pager   # active (running)\n"
  "ss -tlnp | grep crm-analytics               # confirm it is listening\n"
  "journalctl -u crm-analytics -n 20 --no-pager\n\n"
  "analytics_assessment.txt:\n"
  "- Feature responds within acceptable latency; CPU/memory footprint reasonable under test load.\n"
  "- Suggested enhancement: add a systemd timer for scheduled report generation, or cache results.\n\n"
  "Award the mark for a correct .deb install with dependency resolution (apt install ./file.deb, or dpkg -i + "
  "apt -f install), enabling/starting under systemd, verifying active + listening, reading journald logs, and a "
  "sensible assessment. (Lab 11 — Packages; Lab 12 — systemd.)"),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    for label, crit, prompt, cap, pts in PRACTICAL:
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
