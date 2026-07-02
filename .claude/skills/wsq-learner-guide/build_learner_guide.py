#!/usr/bin/env python3
"""build_learner_guide.py — WSQ Learner Guide for TGS-2024048316.

Emits BOTH from the single source (course_content.py):
  * LEARNER-GUIDE.md              (repo root — Markdown mirror)
  * courseware/LG-...docx         (Word: cover, version record, TOC, per-lab steps)

House format (wsq-learner-guide skill): How to Use This Guide, one Heading-1 section
per exam domain, one Heading-2 section per lab (Objective · Goal · What you'll build ·
Step-by-step with exact commands · Test it), a Quick Command Reference, and Support +
assessment flow at the end.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import prodoc
from prodoc import BRAND, DARK, GREY
from course_content import (COURSE, DOMAINS, LABS_BY_NUM, domain_labs, REPO, COURSEWARE)

DOCX_OUT = os.path.join(COURSEWARE, "LG-CompTIA-Linux-Plus-XK0-006.docx")
MD_OUT = os.path.join(REPO, "LEARNER-GUIDE.md")

CODE_BG = "F4F5F7"


# ------------------------------------------------------------------ DOCX helpers
def dpara(doc, text, size=11, bold=False, color=DARK, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    r.font.name = "Arial"; r.font.color.rgb = color
    return p


def dbullet(doc, text, size=10.5):
    return dpara(doc, text, size=size, style="List Bullet")


def dcode(doc, text):
    """Render a shell/code block as shaded monospace paragraphs."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), CODE_BG); pPr.append(shd)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x11, 0x18, 0x27)


def dlabel(doc, label, text, color=BRAND):
    p = doc.add_paragraph()
    r = p.add_run(label + " "); r.bold = True; r.font.size = Pt(10.5); r.font.name = "Arial"; r.font.color.rgb = color
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.name = "Arial"; r2.font.color.rgb = DARK
    return p


# ------------------------------------------------------------------ MD helpers
class MD:
    def __init__(self):
        self.buf = []

    def line(self, s=""):
        self.buf.append(s)

    def code(self, text):
        self.buf.append("```bash")
        self.buf.append(text)
        self.buf.append("```")
        self.buf.append("")

    def text(self):
        return "\n".join(self.buf) + "\n"


def build():
    md = MD()
    doc = Document()
    prodoc.style_headings(doc)
    prodoc.add_cover_page(doc, "Learner Guide", COURSE["title"], COURSE["version"],
                          org_logo=COURSE["org_logo"], course_logo=COURSE["course_logo"])
    prodoc.add_version_control(doc, [
        ("1.0", "1 Jul 2026", "Initial release — aligned to CompTIA Linux+ XK0-006 V8 "
         "exam domains and the 30 hands-on labs.", COURSE["trainer"]),
    ])
    prodoc.add_toc(doc)

    # ---- MD front matter
    md.line(f"# {COURSE['title']} — Learner Guide")
    md.line()
    md.line(f"> **Course code:** {COURSE['code']}  ·  **Exam:** CompTIA Linux+ {COURSE['exam']}  "
            f"·  **Version:** {COURSE['version']}")
    md.line(f"> **Conducted by:** {COURSE['org']} (UEN {COURSE['uen']})")
    md.line()

    # ==================================================== How to Use This Guide
    doc.add_heading("How to Use This Guide", level=1)
    md.line("## How to Use This Guide")
    md.line()
    intro = (f"This Learner Guide accompanies {COURSE['title']} ({COURSE['code']}). It is "
             f"organised by the five CompTIA Linux+ {COURSE['exam']} exam domains and contains "
             "all 30 hands-on labs, each with a goal, the commands you run, and a way to verify "
             "your work. Every lab runs on the free Killercoda Ubuntu Playground — no local "
             "install, virtual machine or credit card is required.")
    dpara(doc, intro); md.line(intro); md.line()
    for b in [
        f"Open the Killercoda Ubuntu Playground: {COURSE['killercoda']}",
        "Work through the labs in order; each maps to one exam sub-objective.",
        "Reset the playground between labs that change kernel, firewall or systemd state.",
        "Type the commands yourself rather than copy-pasting — muscle memory helps in the exam.",
        f"Download the slides and this guide, and take the assessment, on the LMS: {COURSE['lms']}",
    ]:
        dbullet(doc, b); md.line(f"- {b}")
    md.line()

    # what learners need
    dpara(doc, "What you need before starting:", bold=True)
    md.line("**What you need before starting:**"); md.line()
    for b in ["A laptop with a modern web browser and internet access.",
              "Basic comfort with a command line (the labs teach the rest).",
              "A Tertiary Infotech LMS account for courseware and the assessment."]:
        dbullet(doc, b); md.line(f"- {b}")
    md.line()

    # ==================================================== Domains + labs
    for d in DOMAINS:
        doc.add_heading(f"Domain {d['num']} — {d['title']} ({d['weight']}% of the exam)", level=1)
        md.line(f"## Domain {d['num']} — {d['title']} ({d['weight']}% of the exam)")
        md.line()
        # objective coverage line
        objline = "This domain covers exam objectives: " + \
                  "; ".join(f"{o[0]} {o[1]}" for o in d["objs"]) + "."
        dpara(doc, objline, size=10, color=GREY)
        md.line(f"_This domain covers exam objectives:_")
        for o in d["objs"]:
            md.line(f"- **{o[0]}** {o[1]}")
        md.line()

        for lab in domain_labs(d["num"]):
            doc.add_heading(f"Lab {lab['num']} — {lab['title']}", level=2)
            md.line(f"### Lab {lab['num']} — {lab['title']}")
            md.line()

            dlabel(doc, "Exam objective:", f"{lab['objective']} — {lab['obj_title']}")
            dlabel(doc, "Goal:", lab["goal"])
            dlabel(doc, "What you'll build:", lab["build"])
            md.line(f"**Exam objective:** {lab['objective']} — {lab['obj_title']}")
            md.line()
            md.line(f"**Goal:** {lab['goal']}")
            md.line()
            md.line(f"**What you'll build:** {lab['build']}")
            md.line()

            # concepts chips
            concepts = ", ".join(lab["concepts"])
            dpara(doc, "Key concepts: " + concepts, size=10, color=GREY)
            md.line(f"**Key concepts:** {concepts}")
            md.line()

            dpara(doc, "Step-by-step", bold=True, size=11, color=BRAND)
            md.line("**Step-by-step**")
            md.line()
            for i, step in enumerate(lab["steps"], 1):
                dpara(doc, f"Step {i} — {step['title']}", bold=True, size=10.5)
                md.line(f"{i}. **{step['title']}**")
                md.line()
                dcode(doc, step["code"])
                md.code(step["code"])
                dpara(doc, step["explain"], size=10, color=GREY)
                md.line(f"   {step['explain']}")
                md.line()

            dlabel(doc, "Test it / key takeaway:", lab["test"], color=RGBColor(0x0A, 0x7A, 0x3B))
            md.line(f"**Test it / key takeaway:** {lab['test']}")
            md.line()
            doc.add_paragraph()

    # ==================================================== Quick Command Reference
    doc.add_heading("Quick Command Reference", level=1)
    md.line("## Quick Command Reference")
    md.line()
    dpara(doc, "The essential tools and terms per exam domain (see each lab for full usage):")
    md.line("The essential tools and terms per exam domain (see each lab for full usage):")
    md.line()
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t = doc.add_table(rows=0, cols=3); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.add_row().cells
    for i, htxt in enumerate(["Domain", "Lab", "Key commands & concepts"]):
        hdr[i].text = ""; rr = hdr[i].paragraphs[0].add_run(htxt)
        rr.bold = True; rr.font.size = Pt(9.5); rr.font.name = "Arial"
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); prodoc._shade_cell(hdr[i], "C8102E")
    md.line("| Domain | Lab | Key commands & concepts |")
    md.line("|---|---|---|")
    for d in DOMAINS:
        for lab in domain_labs(d["num"]):
            cells = t.add_row().cells
            for j, val in enumerate([str(d["num"]), f"{lab['num']}. {lab['title']}",
                                     ", ".join(lab["concepts"])]):
                cells[j].text = ""; rr = cells[j].paragraphs[0].add_run(val)
                rr.font.size = Pt(9); rr.font.name = "Arial"; rr.font.color.rgb = DARK
            md.line(f"| {d['num']} | {lab['num']}. {lab['title']} | {', '.join(lab['concepts'])} |")
    md.line()

    # ==================================================== Support & assessment flow
    doc.add_heading("Support & Assessment", level=1)
    md.line("## Support & Assessment")
    md.line()
    dpara(doc, "At the assessment step, follow this order:", bold=True)
    md.line("**At the assessment step, follow this order:**"); md.line()
    flow = [
        "TRAQOM — scan the TRAQOM QR code on the LMS and complete the survey.",
        "Assessment Digital Attendance.",
        "Assessment — Written Assessment (SAQ) + Practical Performance (PP), open book.",
        "Submit the assessment answers on the LMS.",
        "Sign the Assessment Summary Record.",
    ]
    for i, s in enumerate(flow, 1):
        dpara(doc, f"{i}. {s}", size=10.5); md.line(f"{i}. {s}")
    md.line()
    contact = (f"Courseware and the assessment are on the LMS — {COURSE['lms']}. "
               "For support: enquiry@tertiaryinfotech.com · +65 6100 0613 · www.tertiarycourses.com.sg")
    dpara(doc, contact, size=10, color=GREY)
    md.line(contact); md.line()
    md.line("---")
    md.line(f"_© 2026 {COURSE['org']}. All rights reserved._")

    prodoc.add_page_numbers(doc)
    prodoc.enable_update_fields(doc)
    doc.save(DOCX_OUT)
    print("wrote", DOCX_OUT)
    # HARD RULE 1 (wsq-learner-guide): deliverables are DOCX + PDF ONLY — no Markdown mirror.
    import subprocess
    try:
        subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir",
                        COURSEWARE, DOCX_OUT], check=True,
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=180)
        print("wrote", DOCX_OUT.replace(".docx", ".pdf"))
    except Exception as e:
        print("PDF render skipped (run soffice manually):", e)
    # Never leave a Markdown mirror in the repo (HARD RULE 1).
    if os.path.exists(MD_OUT):
        os.remove(MD_OUT)


if __name__ == "__main__":
    build()
