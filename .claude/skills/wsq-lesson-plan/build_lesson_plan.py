#!/usr/bin/env python3
"""build_lesson_plan.py — WSQ Lesson Plan (.docx) for TGS-2024048316.

House format (wsq-lesson-plan skill) + HARD RULE 1: the Lesson Plan cites the deck
slide (page) numbers, read from courseware/slide_map.json (written by build_slides.py),
so the trainer can run the class straight from the deck and the two never drift.

Build the slide deck FIRST (build_slides.py) so slide_map.json exists.

Two-day class: final assessment Day 2, 4:00–6:00 PM = 1 hr Written (SAQ) + 1 hr
Practical Performance (PP), open book.
"""
import os, sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
sys.path.insert(0, os.path.join(_HERE, os.pardir, "wsq-learner-guide"))
import prodoc
from prodoc import BRAND, DARK, GREY
from course_content import COURSE, DOMAINS, LEARNING_OUTCOMES, domain_labs, COURSEWARE, LABS_BY_NUM

OUT = os.path.join(COURSEWARE, "LP-CompTIA-Linux-Plus-XK0-006.docx")

# ---- slide map (HARD RULE 1) -------------------------------------------------
_MAP_PATH = os.path.join(COURSEWARE, "slide_map.json")
if not os.path.exists(_MAP_PATH):
    raise SystemExit("slide_map.json not found — run build_slides.py first (HARD RULE 1).")
SMAP = json.load(open(_MAP_PATH))
LAB_PAGE = {int(k): v for k, v in SMAP["labs"].items()}
DOM_PAGE = {int(k): v for k, v in SMAP["domains"].items()}
TOTAL = SMAP["total"]

# per-lab slide span: overview page .. (next lab overview page - 1); last lab .. total
_ordered = sorted(LAB_PAGE.items(), key=lambda kv: kv[1])
LAB_SPAN = {}
for i, (num, pg) in enumerate(_ordered):
    end = (_ordered[i + 1][1] - 1) if i + 1 < len(_ordered) else TOTAL
    LAB_SPAN[num] = (pg, end)


def slides_for(labs):
    """Deck page range covering a list of lab numbers (overview → last test slide)."""
    if not labs:
        return "—"
    a = min(LAB_SPAN[n][0] for n in labs)
    b = max(LAB_SPAN[n][1] for n in labs)
    # extend the start back to the domain section+concepts of the first lab's domain
    first = min(labs)
    dom = LABS_BY_NUM[first]["domain"]
    if DOM_PAGE.get(dom) and first == DOMAINS[dom - 1]["labs"][0]:
        a = DOM_PAGE[dom]
    return f"{a}–{b}"


# ---- colours -----------------------------------------------------------------
HEADER_FILL = "C8102E"; TOPIC_FILL = "FDECEE"; BREAK_FILL = "F2F2F2"
LUNCH_FILL = "FFF4E5"; ASSESS_FILL = "E8F7EE"; ADMIN_FILL = "EFF3FA"
KIND_FILL = {"topic": TOPIC_FILL, "break": BREAK_FILL, "lunch": LUNCH_FILL,
             "assess": ASSESS_FILL, "admin": ADMIN_FILL, "recap": ADMIN_FILL, "lab": None}

# ---- schedule (single source of timing) --------------------------------------
# (start, end, minutes, kind, text, labs)
SCHEDULE = {
 1: ("System Management & Services / User Management", [
    ("9:00", "9:15", 15, "admin", "Welcome, mandatory digital attendance (AM), trainer & learner introduction", None),
    ("9:15", "9:30", 15, "admin", "Course overview, learning outcomes and ground rules", None),
    ("9:30", "11:45", 135, "lab", "Domain 1 — System Management (23%): boot & FHS, kernel modules, LVM storage, "
     "networking, shell operations. Hands-on Labs 1–5", [1, 2, 3, 4, 5]),
    ("11:45", "12:45", 60, "lunch", "Lunch break", None),
    ("12:45", "14:15", 90, "lab", "Domain 1 (cont.): backup & restore, virtualization. Hands-on Labs 6–7", [6, 7]),
    ("14:15", "14:25", 10, "admin", "Digital attendance (PM)", None),
    ("14:25", "17:15", 170, "lab", "Domain 2 — Services and User Management (20%): files, accounts, processes, "
     "packages, systemd, containers. Hands-on Labs 8–13", [8, 9, 10, 11, 12, 13]),
    ("17:15", "17:45", 30, "recap", "Domain 1–2 recap, Q&A and consolidation", None),
    ("17:45", "18:00", 15, "recap", "Day 1 review, preview of Day 2 and PM digital attendance", None),
 ]),
 2: ("Security, Automation, Troubleshooting & Final Assessment", [
    ("9:00", "9:10", 10, "admin", "Mandatory digital attendance (AM) and Day 1 recap", None),
    ("9:10", "11:45", 155, "lab", "Domain 3 — Security (18%): AAA/sudo/PAM, firewalls, OS & account hardening, "
     "cryptography, compliance & audit. Hands-on Labs 14–19", [14, 15, 16, 17, 18, 19]),
    ("11:45", "12:45", 60, "lunch", "Lunch break", None),
    ("12:45", "14:35", 110, "lab", "Domain 4 — Automation, Orchestration & Scripting (17%): Ansible, Bash, Python, "
     "Git, responsible AI. Hands-on Labs 20–24", [20, 21, 22, 23, 24]),
    ("14:35", "14:45", 10, "admin", "Digital attendance (PM)", None),
    ("14:45", "15:35", 50, "lab", "Domain 5 — Troubleshooting (22%): monitoring, storage, network, security, "
     "performance. Hands-on Labs 25–29", [25, 26, 27, 28, 29]),
    ("15:35", "15:55", 20, "lab", "Capstone — end-to-end server build & triage. Hands-on Lab 30", [30]),
    ("15:55", "16:00", 5, "assess", "Revision · Briefing for Assessment · Course feedback & TRAQOM survey", None),
    ("16:00", "17:00", 60, "assess", "Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book", None),
    ("17:00", "18:00", 60, "assess", "Practical Performance (PP) — hands-on Linux tasks, 1 hour, open book. "
     "PM digital attendance; submit answers on the LMS", None),
 ]),
}


def H(text, level=1):
    return doc.add_heading(text, level=level)


def para(text, size=11, bold=False, color=DARK):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Arial"; r.font.bold = bold; r.font.color.rgb = color
    return p


def bullet(text, size=10.5):
    p = doc.add_paragraph(style="List Bullet"); r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Arial"; r.font.color.rgb = DARK
    return p


def set_cell(cell, text, bold=False, size=9.5, color=None, fill=None, align=None):
    cell.text = ""; p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
    if color: r.font.color.rgb = color
    if fill: prodoc._shade_cell(cell, fill)


# ---- build -------------------------------------------------------------------
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc, "LESSON PLAN", COURSE["title"], COURSE["version"],
                      org_logo=COURSE["org_logo"], course_logo=COURSE["course_logo"])
prodoc.add_version_control(doc, [
    (COURSE["version"], "2 Jul 2026", "Initial release — CompTIA Linux+ XK0-006 V8 2-day lesson plan "
     "aligned to the 30 labs, the five exam domains and the slide deck (slide numbers cited).",
     COURSE["trainer"]),
])
prodoc.add_toc(doc)

H("Course Information", 1)
info = [("Course Title", COURSE["title"]),
        ("WSQ Course Reference", COURSE["code"]),
        ("Certification Exam", f"CompTIA Linux+ {COURSE['exam']}"),
        ("Training Provider", f"{COURSE['org']}  (UEN {COURSE['uen']})"),
        ("Duration", "2 days · 8 training hours per day (16 hours)"),
        ("Daily Timing", "9:00 am – 6:00 pm (1-hour lunch; tea breaks within training time)"),
        ("Mode", "Instructor-led, hands-on labs on the free Killercoda Ubuntu Playground"),
        ("Trainer", COURSE["trainer"])]
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
for k, v in info:
    c = t.add_row().cells
    set_cell(c[0], k, bold=True, size=10, fill=TOPIC_FILL); set_cell(c[1], v, size=10)

H("Learning Outcomes", 1)
para("On completion of this course, learners will be able to:")
for lo in LEARNING_OUTCOMES:
    bullet(lo)

H("Assessment", 1)
for a in ["Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book — aligned to the course slides.",
          "Practical Performance (PP) — hands-on Linux tasks, 1 hour, open book — aligned to the course labs.",
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "The final assessment is conducted on Day 2 from 4:00 pm (WA 4:00–5:00, PP 5:00–6:00).",
          "A minimum of 75% attendance is required to be eligible for assessment and funding.",
          f"Courseware and the assessment are on the LMS — {COURSE['lms']}."]:
    bullet(a)

H("Exam Domain Coverage", 1)
para("The course maps 1:1 to the CompTIA Linux+ XK0-006 exam blueprint. Each domain, its exam "
     "weighting, the labs that cover it, and the deck slides:")
dt = doc.add_table(rows=0, cols=4); dt.style = "Table Grid"; dt.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = dt.add_row().cells
for i, h in enumerate(["Domain", "Exam Weight", "Labs", "Deck Slides"]):
    set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for d in DOMAINS:
    c = dt.add_row().cells
    set_cell(c[0], f"{d['num']}. {d['title']}", bold=True, size=9.5, fill=TOPIC_FILL)
    set_cell(c[1], f"{d['weight']}%", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(c[2], f"{d['labs'][0]}–{d['labs'][-1]}", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(c[3], slides_for(d["labs"]), size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

H("Daily Schedule", 1)
para("Each training day runs 9:00 am – 6:00 pm and delivers 8 instructional hours (1-hour lunch; "
     "tea breaks counted within training time). The Slides column cites the deck page range the "
     "trainer runs for each teaching session.")
for day, (theme, rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}", 2)
    tbl = doc.add_table(rows=0, cols=5); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.add_row().cells
    for i, h in enumerate(["Time", "Duration", "Topic / Activity", "Method", "Slides"]):
        set_cell(hdr[i], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
    training = 0
    for start, end, mins, kind, text, labs in rows:
        cells = tbl.add_row().cells; fill = KIND_FILL.get(kind)
        method = {"lab": "Lecture + Hands-on Lab", "topic": "Lecture", "admin": "Admin",
                  "recap": "Discussion", "break": "Break", "lunch": "Break",
                  "assess": "Assessment"}.get(kind, "")
        dur = f"{mins} min" if mins else "—"
        set_cell(cells[0], f"{start}–{end}", bold=(kind == "assess"), size=9.5, fill=fill)
        set_cell(cells[1], dur, size=9.5, fill=fill)
        set_cell(cells[2], text, bold=(kind == "assess"), size=9.5, fill=fill)
        set_cell(cells[3], method, size=9.5, fill=fill)
        set_cell(cells[4], slides_for(labs), size=9.5, fill=fill, align=WD_ALIGN_PARAGRAPH.CENTER)
        if kind != "lunch":
            training += mins
    for row in tbl.rows:
        row.cells[0].width = Inches(1.05); row.cells[1].width = Inches(0.75)
        row.cells[2].width = Inches(3.7); row.cells[3].width = Inches(1.25); row.cells[4].width = Inches(0.85)
    p = doc.add_paragraph()
    r = p.add_run(f"Total training time: {training} minutes ({training // 60} hours).")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY; r.font.name = "Arial"
    assert training == 480, f"Day {day} training minutes = {training}, expected 480"

H("Topic-by-Topic Breakdown", 1)
para("Each lab below cites the deck slide of its overview so the trainer can jump straight to it.")
for d in DOMAINS:
    H(f"Domain {d['num']} — {d['title']} ({d['weight']}%)  ·  Slides {slides_for(d['labs'])}", 2)
    for lab in domain_labs(d["num"]):
        p = doc.add_paragraph()
        r = p.add_run(f"Lab {lab['num']} — {lab['title']}  [Obj {lab['objective']}]  ·  Slide {LAB_PAGE[lab['num']]}")
        r.bold = True; r.font.size = Pt(10.5); r.font.name = "Arial"; r.font.color.rgb = BRAND
        para(lab["goal"], size=10)

H("Resources Required", 1)
for res in ["A laptop with a modern web browser and reliable internet access.",
            f"Free Killercoda Ubuntu Playground — {COURSE['killercoda']} (no local install required).",
            "Course slide deck and Learner Guide (downloaded from the LMS).",
            f"Tertiary Infotech LMS account — {COURSE['lms']} (courseware download and assessment).",
            f"CompTIA Linux+ {COURSE['exam']} Exam Objectives (reference blueprint)."]:
    bullet(res)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
doc.save(OUT)
print("Saved", OUT)
